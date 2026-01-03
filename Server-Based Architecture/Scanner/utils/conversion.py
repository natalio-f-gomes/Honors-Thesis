import zipfile
import datetime
import os
import json
import re
from io import BytesIO
from docx import Document
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from lxml import etree
import logging

logger = logging.getLogger(__name__)


def emu_to_px(emu, dpi=96):
    # 1 inch = 914400 EMU; px = inches * dpi
    return int(round((emu / 914400) * dpi))


def get_core_properties(doc):
    core = doc.core_properties

    def dt(v):
        if not v:
            return None
        if isinstance(v, datetime.datetime):
            return v.isoformat()
        return str(v)

    return {
        "title": core.title,
        "subject": core.subject,
        "creator": core.author,
        "last_modified_by": core.last_modified_by,
        "revision": core.revision,
        "created": dt(core.created),
        "modified": dt(core.modified),
        "category": core.category,
        "comments": core.comments,
        "keywords": core.keywords,
        "content_status": core.content_status,
        "identifier": core.identifier,
        "language": core.language,
        "version": core.version
    }


def get_app_properties(docx_zip):
    # app.xml has things like application, total time, pages, words, etc.
    try:
        xml = docx_zip.read("docProps/app.xml")
        tree = etree.fromstring(xml)
        ns = {"app": "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties",
              "vt": "http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes"}

        def text(tag):
            el = tree.find(f"app:{tag}", ns)
            return el.text if el is not None else None

        return {
            "application": text("Application"),
            "doc_security": text("DocSecurity"),
            "lines": text("Lines"),
            "paragraphs": text("Paragraphs"),
            "words": text("Words"),
            "characters": text("Characters"),
            "company": text("Company"),
            "pages": text("Pages"),
            "total_time": text("TotalTime")
        }
    except KeyError:
        return {}


def extract_paragraphs(doc):
    data = []
    for p in doc.paragraphs:
        # Determine if this paragraph is part of a list
        # Heuristic: inspect numPr
        p_elm = p._p
        numPr = p_elm.find(qn('w:numPr'))
        list_info = None
        if numPr is not None:
            ilvl = numPr.find(qn('w:ilvl'))
            numId = numPr.find(qn('w:numId'))
            list_info = {
                "level": int(ilvl.get(qn('w:val'))) if ilvl is not None else 0,
                "num_id": int(numId.get(qn('w:val'))) if numId is not None else None
            }
        runs = []
        for r in p.runs:
            r_info = {
                "text": r.text,
                "bold": r.bold,
                "italic": r.italic,
                "underline": r.underline,
                "style": r.style.name if r.style else None
            }
            # Hyperlink detection (search ancestor)
            parent = r._element
            hyperlink = None
            while parent is not None:
                if parent.tag == qn('w:hyperlink'):
                    hyperlink = parent.get(qn('r:id'))
                    break
                parent = parent.getparent()
            if hyperlink and hyperlink in p.part.rels:
                target = p.part.rels[hyperlink].target_ref
                r_info["hyperlink"] = target
            runs.append(r_info)
        data.append({
            "text": p.text,
            "style": p.style.name if p.style else None,
            "is_heading": bool(re.match(r"Heading \d+", p.style.name if p.style else "")),
            "heading_level": int(p.style.name.split()[-1]) if p.style and re.match(r"Heading \d+",
                                                                                   p.style.name) else None,
            "list": list_info,
            "runs": runs
        })
    return data


def extract_tables(doc):
    tables_data = []
    for ti, table in enumerate(doc.tables):
        tinfo = {
            "index": ti,
            "rows": []
        }
        for r in table.rows:
            row_data = []
            for c in r.cells:
                # merging: check gridSpan & vMerge
                tc = c._tc
                # Find tcPr element first, then look for gridSpan and vMerge within it
                tcPr = tc.find(qn('w:tcPr'))
                gridSpan = 1
                vMerge = None

                if tcPr is not None:
                    gridSpan_el = tcPr.find(qn('w:gridSpan'))
                    vMerge_el = tcPr.find(qn('w:vMerge'))

                    if gridSpan_el is not None:
                        gridSpan = int(gridSpan_el.get(qn('w:val')))
                    if vMerge_el is not None:
                        vMerge = vMerge_el.get(qn('w:val'))

                row_data.append({
                    "text": c.text,
                    "grid_span": gridSpan,
                    "v_merge": vMerge
                })
            tinfo["rows"].append(row_data)
        tables_data.append(tinfo)
    return tables_data


def extract_images(doc, output_dir):
    # Skip image extraction as requested - return empty list
    return []


def extract_media(doc):
    media = []
    for rel_id, rel in doc.part.rels.items():
        if rel.reltype not in (RT.IMAGE,):
            # Could be charts, embedded objects, hyperlinks, etc.
            media.append({
                "rel_id": rel_id,
                "type": rel.reltype,
                "target": rel.target_ref
            })
    return media


def extract_comments(docx_zip):
    try:
        xml = docx_zip.read("word/comments.xml")
    except KeyError:
        return []
    tree = etree.fromstring(xml)
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    comments = []
    for c in tree.findall("w:comment", ns):
        cid = c.get(qn("w:id"))
        author = c.get(qn("w:author"))
        date = c.get(qn("w:date"))
        text = "".join(c.itertext())
        comments.append({
            "id": cid,
            "author": author,
            "date": date,
            "text": text.strip()
        })
    return comments


def extract_bookmarks(doc):
    bookmarks = []
    for p in doc.paragraphs:
        # Use find_all() method instead of xpath() for namespace-aware search
        for b in p._p.iter(qn('w:bookmarkStart')):
            name = b.get(qn('w:name'))
            bid = b.get(qn('w:id'))
            if name:  # Only add bookmarks that have names
                bookmarks.append({
                    "id": bid,
                    "name": name,
                    "paragraph_text": p.text
                })
    return bookmarks


def extract_notes(docx_zip, kind="footnotes"):
    filename = f"word/{kind}.xml"
    try:
        xml = docx_zip.read(filename)
    except KeyError:
        return []
    tree = etree.fromstring(xml)
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    notes = []
    for fn in tree.findall("w:" + kind[:-1], ns):
        nid = fn.get(qn("w:id"))
        text = "".join(fn.itertext()).strip()
        if nid and text:
            notes.append({"id": nid, "text": text})
    return notes


def extract_hyperlinks_raw(doc):
    links = []
    for rel_id, rel in doc.part.rels.items():
        if rel.reltype == RT.HYPERLINK:
            links.append({"rel_id": rel_id, "target": rel.target_ref})
    return links


def extract_all(docx_file, include_raw_xml=False):
    doc = Document(docx_file)
    with zipfile.ZipFile(docx_file) as z:
        data = {
            "core_properties": get_core_properties(doc),
            "app_properties": get_app_properties(z),
            "paragraphs": extract_paragraphs(doc),
            "tables": extract_tables(doc),
            "images": extract_images(doc, ""),
            "other_media": extract_media(doc),
            "hyperlinks_raw": extract_hyperlinks_raw(doc),
            "comments": extract_comments(z),
            "bookmarks": extract_bookmarks(doc),
            "footnotes": extract_notes(z, "footnotes"),
            "endnotes": extract_notes(z, "endnotes")
        }
    return data


def extract_text_from_docx(resume_model):
    """
    Extract comprehensive data from a DOCX file stored in a Django model.

    Args:
        resume_model: Django model instance with a resume_file field

    Returns:
        dict: Extracted data as JSON-compatible dictionary, or -1 on error
    """
    logger.info(f"[EXTRACT DOCX] Extracting text for resume ID {resume_model.id}")
    try:
        file_field = resume_model.resume_file
        with file_field.open('rb') as docx_file:
            # Create a BytesIO object from the file content
            docx_content = BytesIO(docx_file.read())

            # Extract all data from the DOCX
            extracted_data = extract_all(docx_content, include_raw_xml=False)

            # Get total text length for logging
            total_text = ""
            for paragraph in extracted_data['paragraphs']:
                total_text += paragraph['text'] + "\n"
            for table in extracted_data['tables']:
                for row in table['rows']:
                    for cell in row:
                        total_text += cell['text'] + " "

            logger.info(f"[EXTRACT DOCX] Total extracted text length: {len(total_text)}")
            logger.debug(f"[EXTRACT DOCX] Found {len(extracted_data['paragraphs'])} paragraphs")
            logger.debug(f"[EXTRACT DOCX] Found {len(extracted_data['tables'])} tables")

            return extracted_data

    except Exception as error:
        logger.error(f"[EXTRACT DOCX] Error extracting DOCX text: {error}")
        return -1